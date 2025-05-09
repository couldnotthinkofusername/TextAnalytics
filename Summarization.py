import docx 
import PyPDF2 as pdf
import pdfplumber as pdl
import logging
import datetime




class TextExtraction:
    def DocumentReader(): #Document Extraction
        doc1 = docx.Document('/Users/t2mac/Downloads/sample-docx-files-sample3.docx')
        text = [] 
        tables = []
    
        for words in doc1.paragraphs:
            text.append(words.text)
        
        for table in doc1.tables:
            tables.append(table.table)
        
        #Metadata
        properties = doc1.core_properties
        
        metadata = {
                'title' : properties.title,
                'author' : properties.author,
                'created' : properties.created,
                'modified' : properties.modified,
                'keywords' : properties.keywords,
                'extraction_date': datetime.datetime.now().strftime("%d-%m-%Y"),
                'doc length': len(text),
                'word_count': sum(len(word.split()) for word in text)
                 }

        cleaned_text = CleanText(' '.join(text))
        return (cleaned_text, metadata)
        

            
           
    def pdfReaderwithPDF2(): #PDF2 for small datasets
        pdf1 = pdf.PdfReader('/Users/t2mac/Documents/Sample Docx and Pdfs/PDF Files/dev-example.pdf')
        text = []

        for page in pdf1.pages:
            text.append(page.extract_text())
        
        metadata = pdf1.metadata

        cleaned_text = CleanText(' '.join(text))
        return (cleaned_text, metadata)
            
    def PdfReaderwithPlumber():   #PDF Extraction for large ones
    
        logging.getLogger('pdfplumber').setLevel(logging.ERROR)
    
        pdf1 = pdl.open('/Users/t2mac/Documents/Sample Docx and Pdfs/PDF Files/sample-pdf-files-sample3.pdf')
        text = []
    
        for page in pdf1.pages:
            ext = page.extract_text()

            if ext:
                lines = ext.split()
        
                for line in lines:
                    if line.strip():
                        text.append(line.strip())
        
        metadata = pdf1.metadata
        
        cleaned_text = CleanText(' '.join(text))
        return (cleaned_text, metadata)
    
import re
def CleanText(text):
        
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'(?<=[a-z]) (?=[A-Z])', '. ', text)
        text = re.sub(r'(?<=[a-z]) (?=[a-z])', ' ', text)
        text = re.sub(r'\.([A-Za-z])', r'. \1', text)
        text = re.sub(r'\s*\.\s*','. ', text)
        text = text.strip()
            
        if not text.endswith('.'):
            text += '.'
                
        tokens = text.split()
        seen = set()
        new_tokens = []
        
        for token in tokens:
            if token not in seen:
                seen.add(token)
                new_tokens.append(token)
        
        return ' '.join(new_tokens)


#D1 = print(TextExtraction.DocumentReader())
#PL = print(TextExtraction.PdfReaderwithPlumber())
#PD2 = print(TextExtraction.pdfReaderwithPDF2())
    
